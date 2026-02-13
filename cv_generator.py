import argparse
import re
import sys
from pathlib import Path

from docx import Document


DEFAULT_TEMPLATE_PATH = "CV_template (1).docx"
DEFAULT_OUTPUT_PATH = "CV_generado.docx"
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def iter_doc_paragraphs(doc):
    """Itera párrafos del cuerpo, tablas, encabezados y pies."""
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph


def get_placeholders(doc):
    """
    Busca variables {{ }} en el documento y devuelve una lista
    sin repetir, manteniendo el orden de aparición.
    """
    ordered = []
    seen = set()

    for paragraph in iter_doc_paragraphs(doc):
        found = PLACEHOLDER_PATTERN.findall(paragraph.text)
        for token in found:
            key = token.strip()
            if key not in seen:
                seen.add(key)
                ordered.append(key)

    return ordered


def replace_in_paragraph(paragraph, values):
    """Reemplaza placeholders en un párrafo preservando formato cuando es posible."""
    for run in paragraph.runs:
        original = run.text
        if not original:
            continue

        updated = original
        for key, value in values.items():
            updated = updated.replace(f"{{{{{key}}}}}", value)

        if updated != original:
            run.text = updated

    # Si un placeholder quedó partido en varios runs, hacemos fallback.
    for key, value in values.items():
        token = f"{{{{{key}}}}}"
        if token in paragraph.text:
            paragraph.text = paragraph.text.replace(token, value)


def replace_placeholders(doc, values):
    """Reemplaza placeholders en todo el documento."""
    for paragraph in iter_doc_paragraphs(doc):
        replace_in_paragraph(paragraph, values)


def parse_args():
    parser = argparse.ArgumentParser(description="Generador de CV a partir de plantilla .docx")
    parser.add_argument("--template", default=DEFAULT_TEMPLATE_PATH, help="Ruta de plantilla .docx")
    parser.add_argument("--output", default=DEFAULT_OUTPUT_PATH, help="Ruta de salida .docx")
    return parser.parse_args()


def main():
    args = parse_args()
    template_path = Path(args.template)
    output_path = Path(args.output)

    if not template_path.exists():
        print(f"Error: no existe la plantilla '{template_path}'.", file=sys.stderr)
        return 1

    try:
        doc = Document(template_path)
    except Exception as exc:
        print(f"Error al abrir la plantilla: {exc}", file=sys.stderr)
        return 1

    placeholders = get_placeholders(doc)
    if not placeholders:
        print("No se encontraron placeholders {{ }} en la plantilla.")
        return 1

    print("=== Completar datos del CV ===\n")

    values = {}
    for ph in placeholders:
        values[ph] = input(f"Ingrese {ph}: ")

    replace_placeholders(doc, values)

    try:
        doc.save(output_path)
    except Exception as exc:
        print(f"Error al guardar el CV: {exc}", file=sys.stderr)
        return 1

    print(f"\nCV generado correctamente: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
